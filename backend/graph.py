from typing_extensions import TypedDict
from llm import model
from prompt import drafting_prompt_template, drafting_chat_prompt
from vector_rag import retriever
from langgraph.graph import StateGraph, START, END
import os
import docx
import json
import re
from transformers import pipeline
from docx.shared import Pt

# using the new ComplianceAgent from compliance_checker.py.
from compliance_checker import ComplianceAgent

# Initialize the toxicity classifier.
toxicity_classifier = pipeline("text-classification", model="unitary/unbiased-toxic-roberta")

# Define our state
class State(TypedDict, total=False):
    flow: str
    previous_sow: str
    query_map: dict
    user_query: str
    additional_context: str
    sow: str            # SOW as a JSON string (or raw text) produced by the drafting agent.
    validated_sow: dict # Parsed and validated SOW data.
    formatted_sow: str  # Filename of the generated DOCX.
    compliance_results: dict  # Results from compliance analysis.
    feedback: str
    error: str
    retryCount: int
    doc_file_path: str

def get_relevant_context(state: State):
    context = retriever.invoke(state['user_query'])
    return { 'additional_context': context, 'retryCount': 0 }

def extract_raw_json(response_text):
    """
    Extracts and parses raw JSON from an AI response wrapped in ```json ... ```
    or returns the raw JSON directly if no wrapping exists.
    """
    # Pattern to capture content between ```json ... ``` or just ```
    pattern = r"```(?:json)?\s*([\s\S]*?)```"

    match = re.search(pattern, response_text.strip())
    if match:
        raw_json_str = match.group(1)
    else:
        raw_json_str = response_text.strip()

    try:
        return json.loads(raw_json_str)
    except json.JSONDecodeError as e:
        print(f"❌ JSON parsing failed: {e}")
        return None


def drafting_agent(state: State):
    if state.get('error'):
        previous_content = state.get('sow', '')
        instruction = (f"Below is the previously generated content: {previous_content} "
                       f"The following errors were detected: {state['error']}. "
                       f"Please revise the content accordingly.")
    else:
        instruction = ""

    if state.get('flow') == 'chat':
        prompt = drafting_chat_prompt.invoke({
            "user_query": state['user_query'],
            "previous_sow": state['previous_sow'],
            "feedback": instruction,
        })
        response = model.invoke(prompt)
        return { 'sow': response.content }
    else: 
        prompt = drafting_prompt_template.invoke({
            "query": state['user_query'],
            "additional_context": state['additional_context'],
            "feedback": instruction,
            **state['query_map']
        })
        response = model.invoke(prompt)
        return { 'sow': response.content }

def compliance_agent(state: State):
    # Use the  ComplianceAgent to analyze the SOW.
    agent = ComplianceAgent()
    try:
        # Attempt to parse the SOW as JSON
        sow_data = json.loads(state['sow'])
    except Exception:
        # If parsing fails, treat the entire text as the content to check.
        sow_data = {"sow_text": state['sow']}
    
    # Generate the compliance report using the new agent.
    report = agent.generate_report(sow_data)
    state['compliance_results'] = report
    
    # If any compliance issues are detected, build a brief error message and set it on the state.
    if (report["compliance_score"] < 80 or
        report["missing_fields"] or
        report["structural_issues"] or
        report["content_issues"] or
        report["language_issues"]):
        error_message = "Compliance issues detected: "
        if report["missing_fields"]:
            error_message += f"Missing fields: {', '.join(report['missing_fields'])}. "
        if report["structural_issues"]:
            error_message += f"Structural issues: {', '.join(report['structural_issues'])}. "
        if report["content_issues"]:
            error_message += f"Content issues: {', '.join(report['content_issues'])}. "
        if report["language_issues"]:
            error_message += f"Language issues: {', '.join(report['language_issues'])}. "
        error_message += f"Risk Level: {report['risk_level']}"
        state['error'] = error_message
    return state

def validate_text(text, threshold=0.75):
    try:
        result = toxicity_classifier(text)[0]
    except Exception as e:
        return text, None
    
    toxic_labels = [
        "toxicity", "severe_toxicity", "obscene", "threat",
        "insult", "identity_attack", "sexual_explicit"
    ]
    if result['label'] in toxic_labels and result['score'] > threshold:
        error_msg = (f"[⚠ TOXIC CONTENT DETECTED] Text validation failed: {text}. "
                     f"Reason: {result['label']} with score {round(result['score']*100, 2)}%")
        print(error_msg)
        return text, error_msg
    return text, None

def validate_sow_data(sow_data):
    validated_data = {}
    errors = {}
    for key, value in sow_data.items():
        if isinstance(value, dict):
            validated_subsection = {}
            subsection_errors = {}
            for subkey, subvalue in value.items():
                valid_text, error = validate_text(subvalue)
                validated_subsection[subkey] = valid_text
                if error:
                    subsection_errors[subkey] = error
            validated_data[key] = validated_subsection
            if subsection_errors:
                errors[key] = subsection_errors
        else:
            valid_text, error = validate_text(value)
            validated_data[key] = valid_text
            if error:
                errors[key] = error
    return validated_data, errors

def extract_json_from_sow(raw_sow: str) -> dict:
    extraction_prompt = (
        '''
        Extract the following fields from the given Statement of Work into a JSON object with these keys: 
        "Project Name", "End Date", "Confidentiality", "Intellectual Property", "Termination", "Project Title", "Start Date", "End Date", "Project Name", "SOW Effective Date","Company Information", "Client", "Agreement Date",
"Client Contact", "Contact", "Services Description", "Deliverables",
"Milestones", "Acceptance", "Personnel and Locations", "Representatives",
"Client Representatives", "Contractor Resources", "Terms & Conditions", "Fees", "Expenses",
"Taxes", "Conversion", "Limitation of Liability", "Service Level Agreement", "Assumptions", "Scope of Work",
"Change Process", "Payment Terms", "Timeline", "Company Name", "Client Name",
''' + raw_sow +
        "\n\nOutput the result as a valid JSON and do not format just return pure json."
    )
    response = model.invoke(extraction_prompt)
    try:
        sow_data = json.loads(response.content)
        return sow_data
    except Exception as e:
        raise ValueError("Failed to extract JSON from SOW content: " + str(e))

def validation_agent(state: State):
    try:
        try:
            sow_data = json.loads(state['sow'])
        except Exception as parse_error:
            sow_data = extract_json_from_sow(state['sow'])

        validated_data, errors = validate_sow_data(sow_data)
        if errors:
            state['error'] = json.dumps(errors)
            return { 'feedback': 'REJECTED', 'retryCount': state['retryCount'] + 1 }
        else:
            state['validated_sow'] = validated_data
            state.pop('error', None)
            return { 'feedback': 'ACCEPTED', 'validated_sow': validated_data }
    except Exception as e:
        state['error'] = str(e)
        return { 'feedback': 'REJECTED', 'retryCount': state['retryCount'] + 1 }

def generate_sow(sow_data, output_filename="Generated_SOW_final.docx"):
    markdown = ''
    newLineChar = '\n\n'
    doc = docx.Document()
    doc.add_heading("STATEMENT OF WORK", 0)
    markdown += f'# STATEMENT OF WORK{newLineChar}'
    doc.add_heading(sow_data["Project Name"], level=1)
    markdown += f'## {sow_data["Project Name"]}\n\n'
    
    intro = (f"This Statement of Work (“SOW”) is entered into as of {sow_data['SOW Effective Date']} by and between "
             f"{sow_data['Company Information']} (“{sow_data['Company Name']}”) and {sow_data['Client']} (“{sow_data['Client Name'] or 'Client'}”) under the provisions of "
             f"that certain Master Services Agreement, dated as of {sow_data['Agreement Date']}, by and between {sow_data['Company Name']} and {sow_data['Client Name'] or 'Client'} (the “Agreement”).")
    doc.add_paragraph(intro)
    markdown += f"{intro}{newLineChar}"
    
    # Contact Information
    # doc.add_heading("Contact Information", level=1)
    # markdown += f"## Contact Information{newLineChar}"
    contact_fields = ["Name", "Title", "Address", "Phone", "Mobile", "Email"]
    
    section_order = [
        "Services Description", "Deliverables", "Milestones", "Acceptance",
        "Personnel and Locations", "Representatives", "Client Representatives",
        "Contractor Resources", "Terms & Conditions", "Fees", "Expenses", "Taxes", "Conversion",
        "Limitation of Liability", "Service Level Agreement", "Assumptions", "Change Process"
    ]
    
    for section in section_order:
        doc.add_heading(section, level=1)
        markdown += f"## {section}{newLineChar}"
        try: 
            cleaned_text = re.sub(r"\*\*(.*?)\*\*", r"\1", sow_data[section])
            doc.add_paragraph(cleaned_text)
            markdown += f"{sow_data[section]}{newLineChar}"
        except Exception as e:
            doc.add_paragraph('')
            markdown += newLineChar 
    
        

    #Current State Analysis section

    if "Current State Analysis" in sow_data:
        doc.add_heading("Current State Analysis", level=1)
        markdown += f"## Current State Analysis{newLineChar}"
        doc.add_paragraph(sow_data["Current State Analysis"])
        markdown += f"{sow_data['Current State Analysis']}{newLineChar}"



    #Gap Analysis section

    if "Gap Analysis" in sow_data:
        doc.add_heading("Gap Analysis", level=1)
        markdown += f"## Gap Analysis{newLineChar}"
        doc.add_paragraph(sow_data["Gap Analysis"])
        markdown += f"{sow_data['Gap Analysis']}{newLineChar}"

    # Signature section.
    doc.add_heading("IN WITNESS WHEREOF", level=1)
    markdown += f"IN WITNESS WHEREOF{newLineChar}"
    doc.add_paragraph("Authorized signatures effective as of the effective date of this SOW.")
    markdown += f"Authorized signatures effective as of the effective date of this SOW.{newLineChar}"
    doc.add_paragraph(f"{sow_data['Client Name'] or 'Client'}  Signature: ________________")
    markdown += f"{sow_data['Client Name'] or 'Client'} Signature: ________________{newLineChar}"
    doc.add_paragraph(f"{sow_data['Company Name'] or ''} Signature: ________________")
    markdown += f"{sow_data['Company Name'] or ''} Signature: ________________{newLineChar}"

    static_folder = os.path.join(os.path.dirname(__file__), 'static')
    if not os.path.exists(static_folder):
        os.makedirs(static_folder)

    output_path = os.path.join(static_folder, output_filename)
    doc.save(output_path)

    print(f"✅ SOW document generated: {output_filename}")
    return { "fileName": output_filename, "formatted_sow_md": markdown }

def formatting_agent(state: State):
    output_file = generate_sow(state['validated_sow'])
    state['doc_file_path'] = output_file['fileName']
    state['formatted_sow'] = output_file['formatted_sow_md']
    return state

def agent_router(state: State):
    if state['retryCount'] > 10:
        return 'SUCCESS'
    # If any error exists (from compliance or validation), loop back to drafting.
    if state.get('error'):
        print(f"error: {state.get('error')}")
        return 'REJECTED'
    if state.get('feedback') == 'ACCEPTED':
        return 'SUCCESS'
    print(f"error: {state.get('error')}")
    return 'REJECTED'

# ---- Build the Graph ----
graph_builder = StateGraph(State)
graph_builder.add_node('get_relevant_context', get_relevant_context)
graph_builder.add_node('drafting_agent', drafting_agent)
graph_builder.add_node('compliance_agent', compliance_agent)
graph_builder.add_node('validation_agent', validation_agent)
graph_builder.add_node('formatting_agent', formatting_agent)

graph_builder.add_edge(START, 'get_relevant_context')
graph_builder.add_edge('get_relevant_context', 'drafting_agent')
graph_builder.add_edge('drafting_agent', 'compliance_agent')
graph_builder.add_edge('compliance_agent', 'validation_agent')

# Route based on validation and compliance feedback.
graph_builder.add_conditional_edges(
    "validation_agent",
    agent_router,
    {
        'SUCCESS': 'formatting_agent',
        'REJECTED': 'drafting_agent'
    }
)

graph_builder.add_edge('formatting_agent', END)

graph_agentor = graph_builder.compile()